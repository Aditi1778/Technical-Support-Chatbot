import logging
import pandas as pd
from typing import TypedDict, List
from langgraph.graph import StateGraph, END
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_text_splitters import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from dotenv import load_dotenv
from docx import Document as DocxDocument
import os
import asyncio
import glob

# Logging setup
logging.basicConfig(
    level=logging.INFO,
    format='[%(asctime)s] [%(levelname)s] - %(message)s',
    handlers=[logging.FileHandler("logs/app.log", encoding='utf-8'), logging.StreamHandler()]
)

logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()
api_key = os.getenv("gemini_api_key")
docx_folder = r"raw_files\docs"  # Folder containing .docx files
whatsapp_chat_path = r"raw_files\chat_with_intent.csv"  # Path to WhatsApp chat CSV
embedding_model = "models/text-embedding-004"

# Validate file paths
if not os.path.exists(docx_folder):
    logger.error(f"Document folder not found at: {docx_folder}")
    raise FileNotFoundError(f"Document folder not found at: {docx_folder}")

if not os.path.exists(whatsapp_chat_path):
    logger.error(f"WhatsApp chat file not found at: {whatsapp_chat_path}")
    raise FileNotFoundError(f"WhatsApp chat file not found at: {whatsapp_chat_path}")

class AgentState(TypedDict):
    docs: list
    whatsapp_messages: List[dict]
    embeddings: any
    error: str  # For error handling

# Parse WhatsApp chat data from CSV
async def parse_whatsapp_chat(state: AgentState):
    logger.info("Parsing WhatsApp chat data from CSV...")
    try:
        # Read CSV file
        df = pd.read_csv(whatsapp_chat_path)
        
        # Ensure required columns exist
        required_columns = ['Date', 'Time', 'Sender', 'Message', 'Keywords', 'Intent']
        if not all(col in df.columns for col in required_columns):
            missing = [col for col in required_columns if col not in df.columns]
            logger.error(f"Missing required columns: {missing}")
            return {"error": f"CSV file missing columns: {missing}"}
        
        # Convert to list of dictionaries
        whatsapp_messages = df[required_columns].to_dict('records')
        
        logger.info(f"Parsed {len(whatsapp_messages)} WhatsApp messages.")
        return {"whatsapp_messages": whatsapp_messages}
    except Exception as e:
        logger.error(f"Error parsing WhatsApp chat CSV: {e}")
        return {"error": f"Failed to parse WhatsApp chat CSV: {e}"}

# Asynchronous function for loading multiple .docx documents
async def load_documents(state: AgentState):
    logger.info("Loading .docx documents...")
    try:
        docs = []
        # Find all .docx files in the specified folder
        docx_files = glob.glob(os.path.join(docx_folder, "*.docx"))
        if not docx_files:
            logger.error(f"No DOCX files found in: {docx_folder}")
            return {"error": f"No DOCX files found in: {docx_folder}"}
        
        for docx_path in docx_files:
            logger.info(f"Processing DOCX file: {docx_path}")
            doc = await asyncio.to_thread(DocxDocument, docx_path)
            full_text = [para.text for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            docx_doc = Document(
                page_content="\n".join(full_text),
                metadata={"source": os.path.basename(docx_path)}
            )
            docs.append(docx_doc)
            logger.info(f"Loaded .docx file {os.path.basename(docx_path)} with {len(full_text)} text segments.")
        
        if not docs:
            logger.error("No documents loaded from .docx files.")
            return {"error": "No documents loaded."}
        
        logger.info(f"Loaded {len(docs)} total documents.")
        return {"docs": docs}
    except Exception as e:
        logger.error(f"Error loading .docx documents: {e}")
        return {"error": f"Failed to load .docx documents: {e}"}

# Split text into chunks
async def split_text(state: AgentState):
    logger.info("Splitting text into chunks...")
    try:
        splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        docs = splitter.split_documents(state["docs"])
        logger.info(f"Split into {len(docs)} documents.")
        return {"docs": docs}
    except Exception as e:
        logger.error(f"Error splitting text: {e}")
        return {"error": f"Failed to split text: {e}"}

# Create embeddings and save to FAISS
async def embed_docs(state: AgentState):
    try:
        logger.info("Creating embeddings...")
        embeddings = GoogleGenerativeAIEmbeddings(model=embedding_model, google_api_key=api_key)
        index = FAISS.from_documents(state["docs"], embeddings)
        index.save_local("hr_vector_store")
        logger.info("Embeddings created and saved to hr_vector_store.")
        return {"embeddings": embeddings}
    except Exception as e:
        logger.error(f"Error creating embeddings: {e}")
        return {"error": f"Failed to create embeddings: {e}"}

# Define the workflow
workflow = StateGraph(AgentState)

# Add nodes
workflow.add_node("parse_whatsapp_chat", parse_whatsapp_chat)
workflow.add_node("load_documents", load_documents)
workflow.add_node("split_text", split_text)
workflow.add_node("embed_docs", embed_docs)

# Define edges
workflow.add_edge("parse_whatsapp_chat", "load_documents")
workflow.add_edge("load_documents", "split_text")
workflow.add_edge("split_text", "embed_docs")
workflow.add_edge("embed_docs", END)

# Set entry point
workflow.set_entry_point("parse_whatsapp_chat")

# Compile the graph
graph = workflow.compile()

async def run_knowledge_base():
    # Initialize state
    initial_state = {
        "docs": [],
        "whatsapp_messages": [],
        "embeddings": None,
        "error": ""
    }
    
    try:
        # Execute the workflow
        final_state = await graph.ainvoke(initial_state)
        
        # Check for errors
        if final_state.get("error"):
            logger.error(f"Workflow failed: {final_state['error']}")
            print(f"Error: {final_state['error']}")
        else:
            logger.info("Knowledge base creation completed successfully.")
            print(f"Parsed {len(final_state['whatsapp_messages'])} WhatsApp messages.")
            print(f"Loaded {len(final_state['docs'])} document chunks.")
            print("Vector store saved to 'hr_vector_store'.")
            
    except Exception as e:
        logger.error(f"Workflow execution failed: {e}")
        print(f"Error: {e}")

if __name__ == "__main__":
    asyncio.run(run_knowledge_base())