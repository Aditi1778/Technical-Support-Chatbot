import logging
import pandas as pd
from typing import Annotated, TypedDict, List
from langgraph.graph import StateGraph, END
from langgraph.graph.message import add_messages
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_core.prompts import ChatPromptTemplate
from langchain_text_splitters import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from dotenv import load_dotenv
from langchain_core.documents import Document
from docx import Document as DocxDocument
import os
import glob
import asyncio

# Logging setup
logging.basicConfig(
    level=logging.INFO,
    format='[%(asctime)s] [%(levelname)s] - %(message)s',
    handlers=[logging.FileHandler("app.log"), logging.StreamHandler()]
)

logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()
api_key = os.getenv("gemini_api_key")
docx_folder = r"raw_files\docs"
whatsapp_chat_path = r"raw_files\chat_with_intent.csv"
embedding_model = "models/text-embedding-004"
llm_model = "gemini-2.0-flash"


# Validate file paths
if not os.path.exists(docx_folder):
    logger.error(f"Document folder not found at: {docx_folder}")
    raise FileNotFoundError(f"Document folder not found at: {docx_folder}")

if not os.path.exists(whatsapp_chat_path):
    logger.error(f"WhatsApp chat file not found at: {whatsapp_chat_path}")
    raise FileNotFoundError(f"WhatsApp chat file not found at: {whatsapp_chat_path}")


# Predefined questions
PREDEFINED_QUESTIONS = [
    "What are the steps labeled as creating a new user?",
    "How to create new vendor?",
    "What are the steps to create new customer?",
    "Where to find queue name from case?", 
    "How to create add service to the branch?",
]

# Agent State
class AgentState(TypedDict):
    pages: list
    docs: list
    embeddings: any
    retriever: any
    llm: any
    qa_chain: any
    messages: Annotated[list, add_messages]
    whatsapp_messages: List[dict]
    error: str

# Node Functions
async def parse_whatsapp_chat(state: AgentState):
    logger.info("Parsing WhatsApp chat data from CSV...")
    try:
        df = pd.read_csv(whatsapp_chat_path)
        required_columns = ['Date', 'Time', 'Sender', 'Message', 'Keywords', 'Intent']
        if not all(col in df.columns for col in required_columns):
            missing = [col for col in required_columns if col not in df.columns]
            logger.error(f"Missing required columns: {missing}")
            return {"error": f"CSV file missing columns: {missing}"}
        whatsapp_messages = df[required_columns].to_dict('records')
        logger.info(f"Parsed {len(whatsapp_messages)} WhatsApp messages.")
        return {"whatsapp_messages": whatsapp_messages}
    except Exception as e:
        logger.error(f"Error parsing WhatsApp chat CSV: {e}")
        return {"error": f"Failed to parse WhatsApp chat CSV: {e}"}

# Asynchronous function for loading multiple .docx documents
async def load_documents(state: AgentState):
    logger.info("Loading .docx documents...")
    try:
        docs = []
        # Find all .docx files in the specified folder
        docx_files = glob.glob(os.path.join(docx_folder, "*.docx"))
        if not docx_files:
            logger.error(f"No DOCX files found in: {docx_folder}")
            return {"error": f"No DOCX files found in: {docx_folder}"}
        
        for docx_path in docx_files:
            logger.info(f"Processing DOCX file: {docx_path}")
            doc = await asyncio.to_thread(DocxDocument, docx_path)
            full_text = [para.text for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            docx_doc = Document(
                page_content="\n".join(full_text),
                metadata={"source": os.path.basename(docx_path)}
            )
            docs.append(docx_doc)
            logger.info(f"Loaded .docx file {os.path.basename(docx_path)} with {len(full_text)} text segments.")
        
        if not docs:
            logger.error("No documents loaded from .docx files.")
            return {"error": "No documents loaded."}
        
        logger.info(f"Loaded {len(docs)} total documents.")
        return {"docs": docs}
    except Exception as e:
        logger.error(f"Error loading .docx documents: {e}")
        return {"error": f"Failed to load .docx documents: {e}"}

async def setup_embeddings(state: AgentState):
    logger.info("Initializing embeddings...")
    try:
        embeddings = GoogleGenerativeAIEmbeddings(model=embedding_model, google_api_key=api_key)
        logger.info("Embeddings initialized.")
        return {"embeddings": embeddings}
    except Exception as e:
        logger.error(f"Error initializing embeddings: {e}")
        return {"error": f"Failed to initialize embeddings: {e}"}

async def setup_retriever(state: AgentState):
    logger.info("Setting up retriever...")
    retriever = FAISS.load_local("hr_vector_store", state["embeddings"], allow_dangerous_deserialization=True).as_retriever(search_kwargs={"k": 4})
    logger.info("Retriever setup completed.")
    return {"retriever": retriever}

async def setup_llm(state: AgentState):
    logger.info("Initializing LLM...")
    llm = ChatGoogleGenerativeAI(model=llm_model, google_api_key=api_key)
    logger.info("LLM initialized.")
    return {"llm": llm}

async def build_chain(state: AgentState):
    logger.info("Building QA chain...")
    memory = ConversationBufferMemory(
        memory_key="chat_history",
        return_messages=True,
        output_key="answer",
        k=5
    )
    prompt_template = ChatPromptTemplate.from_messages([
        ("system", """You are a support assistant. Answer only from the manual content provided as document context, providing step-by-step instructions for procedural questions. Treat steps titled 'create a new user' as relevant to user-related queries, even if they describe branch configuration. Use the provided WhatsApp chat history to understand user context, but do not use external knowledge. Accept queries in Hindi, English, or mixed Hinglish, and respond in the same language as the user's query, keeping responses concise and policy-compliant.

WhatsApp Chat History (for context):
{chat_history}"""),
        ("human", "{question}")
    ])
    qa_chain = ConversationalRetrievalChain.from_llm(
        llm=state["llm"],
        retriever=state["retriever"],
        memory=memory,
        condense_question_prompt=prompt_template,
        return_source_documents=True,
        output_key="answer"
    )
    logger.info("QA chain built successfully.")
    return {"qa_chain": qa_chain}

# async def qa_agent(state: AgentState):
#     messages = state.get("messages", [])
#     if not messages:
#         logger.warning("No messages found in state")
#         return {"messages": messages, "error": "No question provided"}
    
#     question = messages[-1].content
#     logger.info(f"Received user question: {question}")
    
#     whatsapp_messages = state.get("whatsapp_messages", [])
#     chat_history = "\n".join(
#         [f"[{msg['Date']} {msg['Time']}] {msg['Sender']}: {msg['Message']}" 
#          for msg in whatsapp_messages[-5:]]
#     ) if whatsapp_messages else "No chat history available."
    
#     try:
#         response = await state["qa_chain"].ainvoke({
#             "question": question,
#             "chat_history": chat_history
#         })
#         answer = response["answer"]
#         retrieved_docs = response.get("source_documents", [])
#         logger.info(f"Retrieved {len(retrieved_docs)} documents for question: {question}")
#         for i, doc in enumerate(retrieved_docs):
#             logger.info(f"Document {i+1}: {doc.page_content[:200]}...")
#         logger.info(f"Generated answer: {answer}")
#         return {"messages": messages + [SystemMessage(content=answer)]}
#     except Exception as e:
#         logger.error(f"Error in QA chain: {e}")
#         return {"messages": messages, "error": f"Failed to process question: {e}"}

async def qa_agent(state: AgentState):
    messages = state.get("messages", [])
    if not messages:
        logger.warning("No messages found in state")
        return {"messages": messages, "error": "No question provided"}
    
    question = messages[-1].content
    logger.info(f"Received user question: {question}")
    
    whatsapp_messages = state.get("whatsapp_messages", [])
    chat_history = "\n".join(
        [f"[{msg['Date']} {msg['Time']}] {msg['Sender']}: {msg['Message']}" 
         for msg in whatsapp_messages[-5:]]
    ) if whatsapp_messages else "No chat history available."
    
    try:
        response = await state["qa_chain"].ainvoke({
            "question": question,
            "chat_history": chat_history
        })
        answer = response["answer"]
        retrieved_docs = response.get("source_documents", [])
        logger.info(f"Retrieved {len(retrieved_docs)} documents for question: {question}")
        for i, doc in enumerate(retrieved_docs):
            doc_content = doc.page_content[:200].encode('ascii', errors='ignore').decode('ascii')
            logger.info(f"Document {i+1}: {doc_content}...")
        logger.info(f"Generated answer: {answer}")
        return {"messages": messages + [SystemMessage(content=answer)]}
    except Exception as e:
        logger.error(f"Error in QA chain: {e}")
        return {"messages": messages, "error": f"Failed to process question: {e}"}

def should_run_qa_agent(state: AgentState):
    return "qa_agent" if state.get("messages") else END

# Compile LangGraph
def compile_graph():
    builder = StateGraph(AgentState)
    builder.add_node("parse_whatsapp_chat", parse_whatsapp_chat)
    builder.add_node("setup_embeddings", setup_embeddings)
    builder.add_node("setup_retriever", setup_retriever)
    builder.add_node("setup_llm", setup_llm)
    builder.add_node("build_chain", build_chain)
    builder.add_node("qa_agent", qa_agent)

    builder.set_entry_point("parse_whatsapp_chat")
    builder.add_edge("parse_whatsapp_chat", "setup_embeddings")
    builder.add_edge("setup_embeddings", "setup_retriever")
    builder.add_edge("setup_retriever", "setup_llm")
    builder.add_edge("setup_llm", "build_chain")
    builder.add_conditional_edges("build_chain", should_run_qa_agent, {"qa_agent": "qa_agent", END: END})
    builder.add_edge("qa_agent", END)

    return builder.compile()
# Global graph instance
graph = compile_graph()

# Function to display predefined questions
def display_questions():
    print("Available questions:")
    for i, q in enumerate(PREDEFINED_QUESTIONS, 1):
        print(f"{i}. {q}")
    print()

# Main execution
async def run_graph():
    print("Welcome to the WhatsApp Chatbot!")
    print("Please select a question by entering its number, or type your own question.")
    print("Type 'quit' or 'exit' to stop.\n")
    
    for i, question in enumerate(PREDEFINED_QUESTIONS, 1):
        print(f"{i}. {question}")
    print()

    while True:
        user_input = input("Your question or number: ").strip()
        
        if user_input.lower() in ["quit", "exit"]:
            print("Goodbye!")
            break
        
        question = None
        if user_input.isdigit():
            index = int(user_input) - 1
            if 0 <= index < len(PREDEFINED_QUESTIONS):
                question = PREDEFINED_QUESTIONS[index]
                logger.info(f"Selected predefined question: {question}")
            else:
                print(f"Please enter a number between 1 and {len(PREDEFINED_QUESTIONS)}.")
                logger.warning(f"Invalid predefined question number: {user_input}")
                continue
        else:
            question = user_input
        
        if not question or not question.strip():
            print("Please enter a valid question or select a number.")
            logger.warning("Empty or invalid question provided, prompting again.")
            continue
        
        logger.info(f"Processing question: {question}")
        
        state = {
            "messages": [HumanMessage(content=question)]
        }
        logger.info(f"Initial state: {state}")
        
        try:
            result = await graph.ainvoke(state)
            # logger.info(f"Graph result: {result}")
            
            messages = result.get("messages", [])
            error = result.get("error", None)
            
            if error:
                print(f"Error: {error}")
                logger.error(f"Graph returned error: {error}")
            elif messages and len(messages) > 1 and isinstance(messages[-1], SystemMessage):
                answer = messages[-1].content
                print(f"Question: {question}")
                print(f"Answer: {answer}\n")
                logger.info(f"Displayed answer: {answer}")
            else:
                print("No answer generated. Please try again.")
                logger.warning("No valid answer found in graph result.")
                
            print("Select another question or type your own (type 'quit' to exit):")
            for i, q in enumerate(PREDEFINED_QUESTIONS, 1):
                print(f"{i}. {q}")
            print()
                
        except Exception as e:
            print(f"An error occurred: {e}")
            logger.error(f"Error running graph: {e}")
    
    return None

if __name__ == "__main__":
    asyncio.run(run_graph())